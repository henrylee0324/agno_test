from pathlib import Path
from typing import Optional, Dict, Any
from enum import Enum
import json

from agno.tools import Toolkit
from agno.utils.log import logger

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert


class FileFormat(Enum):
    TXT = "txt"
    DOCX = "docx"
    PDF = "pdf"


class FileToolsV2(Toolkit):
    def __init__(
        self,
        base_dir: Optional[Path] = None,
        allowed_formats: Optional[list] = None
    ):
        super().__init__(name="file_tools")
        self.base_dir = base_dir or Path.cwd()
        self.allowed_formats = allowed_formats or [f.value for f in FileFormat]
        
        self.register(self.save_document)
        self.register(self.read_document)

    def save_document(
        self, 
        content: str,
        file_name: str,
        format: str = "txt",
        style: Optional[Dict[str, Any]] = None
    ) -> str:
        """Create and save documents with customizable formatting.

        This function allows you to save text content as TXT, DOCX, or PDF files with optional styling.
        For Word and PDF documents, you can customize font size, text alignment, and margins.
        All operations return a JSON response indicating success or failure.

        Args:
            content (str): The text content to save
            file_name (str): Name for the file (without extension)
            format (str, optional): Output format - "txt", "docx", or "pdf". Defaults to "txt"
            style (Dict[str, Any], optional): Document styling options:
                - font_size (int): Font size in points
                - alignment (str): Text alignment ("left", "center", "right")
                - margin (int): Margin size in inches
                Default style uses 11pt font, left alignment, 1-inch margins

        Returns:
            str: JSON string containing operation result:
        """
        logger.info(f"Starting to save document: {file_name}.{format}")
        
        try:
            if format not in self.allowed_formats:
                raise ValueError(f"Unsupported format: {format}")

            # 設定基本樣式
            style = style or {
                "font_size": 11,
                "alignment": "left",
                "margin": 1
            }
            logger.info(f"Applied style settings: font_size={style['font_size']}, alignment={style['alignment']}, margin={style['margin']}")

            file_path = self.base_dir.joinpath(f"{file_name}.{format}")
            logger.info(f"Target file path: {file_path}")

            # 建立word文件
            if format in [FileFormat.DOCX.value, FileFormat.PDF.value]:
                logger.info("Creating Word document...")
                doc = docx.Document()
                
                # 設定邊距
                logger.info("Applying document margins...")
                for section in doc.sections:
                    margin = Inches(style["margin"])
                    section.top_margin = margin
                    section.bottom_margin = margin
                    section.left_margin = margin
                    section.right_margin = margin

                # 添加內容
                logger.info("Adding content to document...")
                p = doc.add_paragraph(content)
                logger.info("Content: \n" + content)
                
                # 套用格式
                logger.info("Applying text formatting...")
                p.runs[0].font.size = Pt(style["font_size"])
                p.alignment = {
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT
                }.get(style["alignment"], WD_PARAGRAPH_ALIGNMENT.LEFT)

                # 儲存docx
                doc_path = file_path
                if format == FileFormat.PDF.value:
                    doc_path = self.base_dir.joinpath(f"{file_name}_temp.docx")
                    logger.info(f"Creating temporary Word document: {doc_path}")
                
                logger.info(f"Saving Word document to: {doc_path}")
                doc.save(doc_path)

                # 轉換為PDF
                if format == FileFormat.PDF.value:
                    logger.info(f"Starting PDF conversion from: {doc_path}")
                    logger.info(f"Target PDF path: {file_path}")
                    convert(str(doc_path), str(file_path))
                    logger.info("PDF conversion completed successfully")
                    
                    logger.info(f"Removing temporary file: {doc_path}")
                    doc_path.unlink()
                    logger.info("Temporary file removed")

            else:  # txt格式
                logger.info(f"Creating text file: {file_path}")
                file_path.write_text(content)
                logger.info("Text file created successfully")

            logger.info(f"Document successfully saved to: {file_path}")
            return json.dumps({
                "success": True,
                "file_path": str(file_path.absolute())  # 使用絕對路徑
            })

        except Exception as e:
            error_msg = f"Error saving document: {str(e)}"
            logger.error(error_msg)
            return json.dumps({
                "success": False,
                "error": error_msg
            })

    def read_document(
        self, 
        file_path: str
    ) -> str:
        """Extract content from TXT or DOCX files.

        Reads and returns the text content from supported document formats.
        For Word documents, extracts text from all paragraphs and joins them with newlines.

        Args:
            file_path (str): Path to the document file

        Returns:
            str: JSON string containing operation result:
        """
        logger.info(f"Starting to read document: {file_path}")
        
        try:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            content = ""
            if path.suffix == '.docx':
                logger.info("Reading Word document...")
                doc = docx.Document(path)
                content = "\n".join([p.text for p in doc.paragraphs])
                logger.info("Word document successfully read")
            elif path.suffix == '.txt':
                logger.info("Reading text file...")
                content = path.read_text()
                logger.info("Text file successfully read")
            else:
                raise ValueError(f"Unsupported file type: {path.suffix}")

            logger.info(f"Successfully read document from: {file_path}")
            return json.dumps({
                "success": True,
                "content": content
            })

        except Exception as e:
            error_msg = f"Error reading document: {str(e)}"
            logger.error(error_msg)
            return json.dumps({
                "success": False,
                "error": error_msg
            })